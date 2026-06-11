#!/usr/bin/env python3
"""Generate user-guide.docx — PBIRS Power BI Report System (English)."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup (A4) ───────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin  = sec.bottom_margin = Cm(2)

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
GOLD  = RGBColor(0xF0, 0xA5, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0x40, 0x40, 0x40)
GREEN = RGBColor(0x1A, 0x7A, 0x3C)
RED   = RGBColor(0xC0, 0x00, 0x00)
BLUE  = RGBColor(0x00, 0x55, 0xAA)

# ── Helpers ───────────────────────────────────────────────────────────────────
def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name   = "Segoe UI"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color

def shd_cell(cell, fill_hex):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(s)

def shd_para(p, fill_hex):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), fill_hex)
    p._p.get_or_add_pPr().append(s)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    shd_para(p, '003366')
    r = p.add_run(f"  {text}")
    font(r, 14, bold=True, color=WHITE)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font(r, 12, bold=True, color=NAVY)
    bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'F0A500')
    bdr.append(bot); p._p.get_or_add_pPr().append(bdr)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    font(r, 11, bold=True, color=NAVY)

def body(text, indent=0, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    font(r, 10.5, color=color or GRAY)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Cm(0.6)
    shd_para(p, 'FFF8E1')
    r1 = p.add_run("Note:  "); font(r1, 10, bold=True, color=RGBColor(0xB8,0x5C,0x00))
    r2 = p.add_run(text);     font(r2, 10, color=RGBColor(0x5C,0x3A,0x00))

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Cm(0.6)
    shd_para(p, 'E8F5E9')
    r1 = p.add_run("Tip:  "); font(r1, 10, bold=True, color=GREEN)
    r2 = p.add_run(text);     font(r2, 10, color=RGBColor(0x1B,0x5E,0x20))

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Cm(0.6)
    shd_para(p, 'FFEBEE')
    r1 = p.add_run("Important:  "); font(r1, 10, bold=True, color=RED)
    r2 = p.add_run(text);           font(r2, 10, color=RGBColor(0x7F,0x00,0x00))

def step(n, bold_part, rest=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.left_indent = Cm(0.4)
    r1 = p.add_run(f"Step {n}:  "); font(r1, 11, bold=True, color=NAVY)
    r2 = p.add_run(bold_part);      font(r2, 11, bold=True, color=GRAY)
    if rest:
        r3 = p.add_run(f"  {rest}"); font(r3, 11, color=GRAY)

def bullets(items, indent=0.8):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            r1 = p.add_run(item[0]); font(r1, 10.5, bold=True, color=NAVY)
            r2 = p.add_run(item[1]); font(r2, 10.5, color=GRAY)
        else:
            font(p.add_run(item), 10.5, color=GRAY)

def numbered(items, indent=0.8):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(4)
        if isinstance(item, tuple):
            r1 = p.add_run(item[0]); font(r1, 10.5, bold=True, color=NAVY)
            r2 = p.add_run(item[1]); font(r2, 10.5, color=GRAY)
        else:
            font(p.add_run(item), 10.5, color=GRAY)

def img(caption="[INSERT SCREENSHOT HERE]", h=6):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); shd_cell(c, 'E3EEFF')
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(h*8); p.paragraph_format.space_after = Pt(h*8)
    font(p.add_run(f"  {caption}  "), 10, italic=True, color=RGBColor(0x55,0x6B,0xAB))
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    font(p2.add_run(caption.strip('[]')), 9, italic=True, color=RGBColor(0x99,0x99,0x99))

def code_block(lines):
    """Render a code / command block (light grey box)."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        shd_para(p, 'F2F2F2')
        font(p.add_run(line), 9.5, color=RGBColor(0x1A,0x1A,0x1A))
        p._p.get_or_add_pPr().append(_mono_font_pPr())

def _mono_font_pPr():
    # just returns a dummy element — actual mono look from font below
    return OxmlElement('w:jc')  # placeholder, ignored

def sp(n=1):
    for _ in range(n):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)

def hline():
    p = doc.add_paragraph()
    bdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
    b.set(qn('w:space'),'1');    b.set(qn('w:color'),'CCCCCC')
    bdr.append(b); p._p.get_or_add_pPr().append(bdr)

def simple_table(headers, rows, col_fill='003366', alt=('FFFFFF','F5F5F5')):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.cell(0,i); shd_cell(c, col_fill)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(h), 10.5, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        fill = alt[ri%2]
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; shd_cell(c, fill)
            bold = (ci==0); col = NAVY if ci==0 else GRAY
            font(c.paragraphs[0].add_run(val), 10, bold=bold, color=col)
    sp()


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
shd_para(p, '003366')
font(p.add_run("\n  BANCO NACIONAL DE COMERCIO DE TIMOR LESTE  \n"), 16, bold=True, color=WHITE)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
font(p.add_run("POWER BI REPORT SYSTEM"), 22, bold=True, color=NAVY)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("User Guide"), 14, italic=True, color=GOLD)

img("[INSERT BNCTL LOGO HERE]", h=4)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("For end users  ·  No technical knowledge required"), 11, italic=True, color=GRAY)
sp()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Version 1.0  |  Last updated: 2026-06-11"), 9, color=RGBColor(0xAA,0xAA,0xAA))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — ACCESSING & VIEWING REPORTS
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 1 — ACCESSING AND VIEWING REPORTS ON THE SERVER")
sp()
body("All reports are published on the internal Power BI Report Server (PBIRS). "
     "You can access them from any computer on the internal network — no software installation needed.")

h2("1.1  Open the Report Server in a browser")
step(1, "Open a web browser", "(Chrome, Edge, or Firefox).")
step(2, "Type the following address and press Enter:")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
font(p.add_run("http://10.0.40.122/reports/browse/REPORT_V2"), 12, bold=True, color=NAVY)
img("[SCREENSHOT: Browser address bar with URL entered]")
step(3, "Log in if prompted:")
bullets([("Username: ","administrator"), ("Password: ","contact your system administrator")])
img("[SCREENSHOT: Windows / PBIRS login screen]")

h2("1.2  Find and open a report")
step(1, "The home screen shows all published reports in the REPORT_V2 folder.")
img("[SCREENSHOT: PBIRS home screen showing report list]")
step(2, 'Click the report name you want to view (e.g. Credit Report, Accounting, BNCTL_Treasury_Reports).')
img("[SCREENSHOT: Clicking a report name]")
step(3, "The report opens in the browser. Use the slicers (filters) on the side or top to filter data.")
img("[SCREENSHOT: Report open in browser with slicers]", h=8)
note("If the report shows no data or an error, contact your administrator to check the data pipeline.")

h2("1.3  Export a report")
step(1, 'Click the  ···  (More options) button or the Export icon on the report toolbar.')
img("[SCREENSHOT: Export / More options button on report toolbar]")
step(2, "Choose your export format:")
bullets([("PDF — ","recommended for printing or emailing"),
         ("Excel — ","for further analysis"),
         ("PowerPoint — ","for presentations")])
img("[SCREENSHOT: Export format selection menu]")
step(3, "The file downloads to your computer's Downloads folder.")
tip("Use PDF when printing a formal report for management.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — FIRST-TIME SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 2 — FIRST-TIME SETUP (RUN ONCE PER MACHINE)")
sp()
body("This section is for the person who manages or edits reports through the AI assistant. "
     "Run these steps once on each new machine. If you only view reports in a browser, skip this part.")

h2("2.1  Prerequisites")
bullets(["Windows 10/11 machine on the internal network.",
         "Power BI Desktop RS installed (see 2.3 if not).",
         "Git for Windows installed (setup script will install it if missing)."])

h2("2.2  Clone the repository")
body("The repository contains all report source files. Clone it once.")

# WSL / Windows split
h3("Option A — Windows (PowerShell or Command Prompt)")
body("Open PowerShell as Administrator and run:", indent=0.6)
code_block([
    "git clone https://github.com/phamthanhhai003/pbirs-report",
    "cd pbirs-report",
])
img("[SCREENSHOT: PowerShell with clone command]")

h3("Option B — WSL (Windows Subsystem for Linux)")
body("Open your WSL terminal and run:", indent=0.6)
code_block([
    "git clone https://github.com/phamthanhhai003/pbirs-report /mnt/d/pbirs-report",
    "cd /mnt/d/pbirs-report",
])
note("The repo must be on your Windows drive (D:\\) so that PowerShell scripts can access it directly. "
     "Cloning into a Linux-only path (e.g. /home/...) will break the scripts.")
img("[SCREENSHOT: WSL terminal with clone command]")

h2("2.3  Run the setup script")
body("The setup script installs all dependencies and configures the machine.")

h3("Option A — Windows (PowerShell, run as Administrator)")
code_block([
    "cd D:\\pbirs-report",
    "powershell -ExecutionPolicy Bypass -File scripts\\setup.ps1",
])

h3("Option B — WSL")
code_block([
    "cd /mnt/d/pbirs-report",
    "powershell.exe -ExecutionPolicy Bypass -File scripts/setup.ps1",
])

sp()
body("The script automatically installs:")
simple_table(
    ["Step", "What it installs / does"],
    [
        ("1", "Git for Windows — version control"),
        ("2", "Node.js LTS — required for Claude Code"),
        ("3", "Claude Code CLI — the AI assistant"),
        ("4", "Tabular Editor 2 — DAX read/write engine"),
        ("5", "config.ps1 — server URL, credentials (pre-filled for dev)"),
        ("6", "Git hooks — wires up automatic DAX extraction on commit"),
    ]
)
img("[SCREENSHOT: Setup script running in terminal]")

h2("2.4  Install Power BI Desktop RS (if not already installed)")
body("Power BI Desktop RS cannot be installed silently by the script. If it is not present:")
numbered([
    "Download from Microsoft: search 'Power BI Report Server download' on Microsoft's website.",
    "Run the installer and follow the on-screen steps.",
    'Choose "Power BI Desktop (optimized for Power BI Report Server)" during installation.',
    "After installation, re-run setup.ps1 (step 2.3) to complete configuration.",
])
img("[SCREENSHOT: Power BI Desktop RS installer]")
note("Make sure you install Power BI Desktop RS, NOT the regular Power BI Desktop from the Microsoft Store. "
     "They are different applications.")

h2("2.5  Start the AI assistant")
body("After setup, start Claude Code from the repository folder.")

h3("Option A — Windows (PowerShell)")
code_block(["cd D:\\pbirs-report", "claude"])

h3("Option B — WSL")
code_block(["cd /mnt/d/pbirs-report", "claude"])

img("[SCREENSHOT: Claude Code starting in terminal]")
tip("You can create a shortcut or alias so you don't have to type the path each time.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — REQUESTING CHANGES VIA AI
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 3 — REQUESTING REPORT CHANGES THROUGH THE AI")
sp()
body("When you need to change what is displayed in a report — hide/show a metric, change a colour, "
     "adjust a value — describe the change in plain English. No technical knowledge required.")

h2("3.1  Before each session")
body("Make sure these two conditions are met before giving any instruction:")
bullets([
    "Power BI Desktop RS is open with the correct .pbix file loaded.",
    "The window title bar shows the file name (e.g. Accounting - Power BI Desktop).",
])
img("[SCREENSHOT: Power BI Desktop RS title bar showing file name]")
warn("If Power BI Desktop RS is not open, the AI will refuse to make changes and prompt you to open the file first.")

h2("3.2  How to give instructions")
body("Open the AI chat and describe the change in plain English. Examples of valid instructions:")
bullets([
    '"Remove the Total Write-Off card from Extra Accountable"',
    '"Change the Liquidity Ratio colour from red to yellow"',
    '"Hide the Total Outstanding Loan metric in the Provision report"',
    '"Add back the Total Loans card that was removed earlier"',
])
img("[SCREENSHOT: AI chat interface with instruction typed]")
note("If you have two report files open at the same time, always include the file name in your instruction "
     "so the AI targets the correct window.")

h2("3.3  What happens after you send an instruction")
numbered([
    ("AI processes:  ","finds the correct metric and applies the change in Power BI Desktop RS."),
    ("Auto-save:  ","the system automatically saves the .pbix file after the change."),
    ("AI notifies:  ",'shows "Change applied and saved to Desktop. Open report to verify."'),
])
img("[SCREENSHOT: AI chat showing change applied notification]")

h2("3.4  Verify the result in Power BI Desktop RS")
step(1, "Switch to Power BI Desktop RS", "(click it on the taskbar).")
img("[SCREENSHOT: Power BI Desktop RS icon on taskbar]")
step(2, "Navigate to the page containing the changed metric.")
step(3, "Check that the result looks correct.")
img("[SCREENSHOT: Power BI Desktop RS showing updated report]", h=8)

h2("3.5  Choose what to do next — the 3 options")
body("After verifying the result, the AI presents exactly three choices:")

t = doc.add_table(rows=4, cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(["Choice","What it does"]):
    c=t.cell(0,i); shd_cell(c,'003366')
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(h),10.5,bold=True,color=WHITE)
opts = [
    ("1 — Approve and push to server",
     "Saves change history (Git) and updates the live report on the server immediately.\n"
     "Everyone viewing the report on PBIRS will see the new version within seconds."),
    ("2 — Keep editing",
     "Describe another change. AI applies it and shows the 3 choices again.\n"
     "Use this to make multiple changes before pushing once."),
    ("3 — Revert",
     "Undoes the change. Report is restored to the last saved state in the repo.\n"
     "Use when the result is not what you expected."),
]
fills = ['E8F5E9','E3F2FD','FFF3E0']
for i,(opt,desc) in enumerate(opts):
    r=t.rows[i+1]
    for j,txt in enumerate([opt,desc]):
        c=r.cells[j]; shd_cell(c,fills[i])
        font(c.paragraphs[0].add_run(txt),10.5,bold=(j==0),color=NAVY if j==0 else GRAY)
sp()
img("[SCREENSHOT: AI chat showing the 3 options]")

h2("3.6  After choosing Option 1 — Approve and push")
body("The system automatically performs 3 steps:")
numbered([
    ("Save history (Git commit):  ","records the exact change with a timestamp. Can be reviewed or reversed at any time."),
    ("Push to GitHub:  ","syncs the source code to the remote repository."),
    ("Upload to PBIRS:  ","pushes the updated .pbix to http://10.0.40.122 — report is live immediately."),
])
img("[SCREENSHOT: AI chat showing commit + push + upload success]")
tip("After the AI reports success, press F5 on the PBIRS browser page to see the updated report.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — COMMON SCENARIOS
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 4 — COMMON SCENARIOS")
sp()

h2('4.1  Hide a card from a report')
body('Example: Remove the "Total Write-Off" card from the Extra Accountable report.')
step(1, 'Tell the AI:', '"Remove the Total Write-Off card from Extra Accountable"')
step(2, 'AI applies the change → verify in Power BI Desktop RS.')
step(3, 'Looks correct → Choose 1 to push to server.')
img("[SCREENSHOT: Before and after hiding the card]")

h2('4.2  Restore a card that was previously hidden')
body('Example: Bring back the "Total Write-Off" card.')
step(1, 'Tell the AI:', '"Add back the Total Write-Off card in Extra Accountable"')
step(2, 'AI restores from the saved repo state → verify → Choose 1.')
img("[SCREENSHOT: Card restored in the report]")

h2('4.3  Change the colour of a metric')
body('Example: Liquidity Ratio is showing in red — change it to yellow.')
step(1, 'Tell the AI:', '"Change the Liquidity Ratio colour to yellow"')
step(2, 'AI updates the colour logic → verify → Choose 1.')
img("[SCREENSHOT: Liquidity Ratio displaying in yellow]")

h2('4.4  Make several changes before pushing (use Choice 2)')
body('When you need to make multiple edits before updating the server:')
numbered([
    'Tell the AI the first change → AI edits → verify → Choose 2.',
    'Tell the AI the next change → AI edits → verify → Choose 2.',
    'Repeat until all changes are done → Choose 1 to push everything at once.',
])
note("All changes are bundled into a single server update — faster and cleaner history.")

h2('4.5  Undo a change (Choose 3 — Revert)')
step(1, 'Verify the result in Power BI Desktop RS → something looks wrong.')
step(2, 'Go back to the AI chat → Choose 3.')
step(3, 'AI restores the previous state → verify again.')
img("[SCREENSHOT: Choosing option 3 in the AI chat]")

h2('4.6  Two report files open at the same time')
body('Always include the file name in your instruction:')
simple_table(
    ["You say", "AI targets"],
    [
        ('"Remove Total NPL from Credit Report"', "Credit Report window"),
        ('"Change colour in Accounting"',          "Accounting window"),
        ('"Remove Total NPL"',                     "AI asks which file — slower"),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 5 — TROUBLESHOOTING")
sp()

problems = [
    ('"PBI Desktop RS not running" error',
     "Power BI Desktop RS is not open or no file is loaded.",
     ["Open Power BI Desktop RS.",
      "Open the .pbix file you want to edit.",
      "Wait for the file to fully load (title bar shows the file name).",
      "Retry your instruction."]),
    ("Change applied but nothing changed in Desktop",
     "Multiple .pbix files are open — AI targeted the wrong window.",
     ['Close files you are not editing.',
      'Or specify the file name: "Remove card X from Credit Report".']),
    ("Cannot access the report server URL",
     "The machine is not on the internal network.",
     ["Check your network connection.",
      "Make sure you are on LAN or VPN — mobile data will not work.",
      "Try pinging 10.0.40.122."]),
    ("Report shows old data after an update",
     "The browser is serving a cached version.",
     ["Press F5 or Ctrl+F5 to hard-refresh.",
      "If still outdated, contact IT to check the data pipeline."]),
    ("Setup script fails or cannot find a dependency",
     "Missing prerequisite or network issue during download.",
     ["Run the script again — it skips steps that already completed.",
      "Check internet connectivity.",
      "Run PowerShell as Administrator.",
      "Contact IT if the error persists."]),
    ("WSL: powershell.exe not found",
     "WSL cannot locate the Windows PowerShell executable.",
     ['Check that PowerShell is installed on Windows.',
      'In WSL, run: which powershell.exe  — it should return /mnt/c/Windows/System32/WindowsPowerShell/v1.0/powershell.exe',
      'If not found, add Windows system32 to WSL PATH.']),
]

for title, cause, fixes in problems:
    h3(f"  {title}")
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r1=p.add_run("Cause:  "); font(r1,10.5,bold=True,color=RED)
    r2=p.add_run(cause);      font(r2,10.5,color=GRAY)
    p2=doc.add_paragraph(); p2.paragraph_format.left_indent=Cm(0.4)
    r3=p2.add_run("Fix:"); font(r3,10.5,bold=True,color=GREEN)
    for f in fixes:
        bp=doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent=Cm(1.2); bp.paragraph_format.space_after=Pt(3)
        font(bp.add_run(f),10.5,color=GRAY)
    sp()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PART 6 — REPORT LIST
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 6 — AVAILABLE REPORTS ON THE SERVER")
sp()

h2("6.1  Credit Report")
simple_table(
    ["Report Name", "Server Path"],
    [("Provision Report",         "/REPORT_V2/Credit Report"),
     ("Provision v2",             "/REPORT_V2/Credit Report"),
     ("Repayment Report",         "/REPORT_V2/Credit Report"),
     ("Disbursement Summary",     "/REPORT_V2/Credit Report"),
     ("Disbursement Monthly",     "/REPORT_V2/Credit Report"),
     ("Loan Sector Daily",        "/REPORT_V2/Credit Report"),
     ("Loan Sector Yearly",       "/REPORT_V2/Credit Report"),
     ("Extra Accountable Report", "/REPORT_V2/Credit Report")]
)

h2("6.2  Treasury & Accounting Reports")
simple_table(
    ["Report Name", "Server Path"],
    [("Liquidity Report",             "/REPORT_V2/BNCTL_Treasury_Reports"),
     ("Remittance Daily",             "/REPORT_V2/BNCTL_Treasury_Reports"),
     ("Remittance Quarterly",         "/REPORT_V2/BNCTL_Treasury_Reports"),
     ("Accounting (Assets / BS / P&L)","/REPORT_V2/Accounting")]
)

# ══════════════════════════════════════════════════════════════════════════════
# PART 7 — SYSTEM INFO & CONTACTS
# ══════════════════════════════════════════════════════════════════════════════
h1("PART 7 — SYSTEM INFORMATION & SUPPORT CONTACTS")
sp()
simple_table(
    ["Item", "Value"],
    [("PBIRS Server URL",    "http://10.0.40.122/reports/browse/REPORT_V2"),
     ("Git Repository",      "https://github.com/phamthanhhai003/pbirs-report"),
     ("Power BI Desktop RS", r"C:\Program Files\Microsoft Power BI Desktop RS"),
     ("Tabular Editor 2",    r"C:\Program Files (x86)\Tabular Editor"),
     ("Repo root (Windows)", r"D:\pbirs-report"),
     ("Repo root (WSL)",     "/mnt/d/pbirs-report"),
     ("IT Support contact",  "[FILL IN HERE]"),
     ("IT Support email",    "[FILL IN HERE]")]
)

sp(2); hline()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("BANCO NACIONAL DE COMERCIO DE TIMOR LESTE  —  Internal Document  —  v1.0  —  2026-06-11"),
     9, italic=True, color=RGBColor(0xAA,0xAA,0xAA))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/mnt/d/pbirs-report/docs/user-guide.docx"
doc.save(out)
print(f"Saved: {out}")
